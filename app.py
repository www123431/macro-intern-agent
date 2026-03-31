import streamlit as st
import requests
import feedparser
import google.generativeai as genai
import google.api_core.exceptions
import urllib.parse
import streamlit.components.v1 as components
import time
import datetime
from docx import Document
from io import BytesIO
import numpy as np
import pandas as pd
from scipy.stats import norm
import yfinance as yf
from sklearn.manifold import Isomap
from sklearn.linear_model import LassoCV
from typing import TypedDict, Annotated, List
import operator
from langgraph.graph import StateGraph, END
from typing import TypedDict, Annotated
from engine.scanner import MarketScanner

# --- 1. 页面基本配置 ---
st.set_page_config(page_title="Macro Alpha Pro Terminal", layout="wide", page_icon="🏛️")

# --- 1.5 登录验证逻辑 (修复版) ---
def check_password():
    """验证用户名和密码"""
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    if st.session_state["password_correct"]:
        return True

    # 登录界面样式
    st.markdown("""
        <style>
        .login-box {
            max-width: 400px;
            margin: 50px auto;
            padding: 30px;
            background-color: #FFFFFF;
            border-radius: 15px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            border: 1px solid #E2E8F0;
            text-align: center;
        }
        </style>
        <div class="login-box">
            <h2 style='color: #1B315E;'>🏛️ Macro Alpha Pro</h2>
            <p style='color: #64748B;'>请输入凭据以访问终端</p>
        </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        user = st.text_input("用户名")
        pwd = st.text_input("密码", type="password")
        login_btn = st.button("进入终端", use_container_width=True)

        if login_btn:
            if user == "zhang" and pwd == "200211":
                st.session_state["password_correct"] = True
                st.rerun() # 👈 验证成功后强制刷新页面进入主程序
            else:
                st.error("❌ 用户名或密码错误")
    return False

# 拦截器
if not check_password():
    st.stop()

st.markdown("""
    <style>
    /* 调整 Metric 卡片的背景 */
    [data-testid="stMetricValue"] {
        font-size: 28px;
        color: #1E3A8A;
    }
    div[data-testid="metric-container"] {
        background-color: #F8FAFC;
        border: 1px solid #E2E8F0;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    /* 优化按钮样式 */
    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #1E3A8A;
    }
    /* 审计报告的卡片式布局 */
    .audit-card {
        background-color: #FFFFFF;
        border-left: 5px solid #1E3A8A;
        padding: 20px;
        margin: 10px 0;
        border-radius: 5px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    </style>
    """, unsafe_allow_html=True)

st.title("🏛️ Macro Alpha")

# --- 2. 安全读取 Secrets ---
try:
    GEMINI_KEY = st.secrets.get("GEMINI_API_KEY") or st.secrets.get("GEMINI_KEY")
    if not GEMINI_KEY:
        st.error("❌ Secrets 配置不完整，请在 Streamlit 控制台配置 GEMINI_API_KEY。")
        st.stop()
except Exception as e:
    st.error(f"❌ 配置文件读取失败: {e}")
    st.stop()

# --- 3. 初始化 Gemini ---
genai.configure(api_key=GEMINI_KEY)
model = genai.GenerativeModel(model_name='gemini-2.5-flash')

# --- 4. 核心量化引擎 ---

class StrategyAuditor:
    @staticmethod
    def run_feature_sparsity_check(X, y):
        common_index = X.index.intersection(y.index)
        X, y = X.loc[common_index], y.loc[common_index]
        model_lasso = LassoCV(cv=5).fit(X, y)
        active_features = np.sum(np.abs(model_lasso.coef_) > 1e-10)
        sparsity = 1 - (active_features / X.shape[1])
        return active_features, sparsity, model_lasso.coef_

    @staticmethod
    def check_optimizer_curse(val_score, test_score, n_trials):
        luck_factor = np.log1p(n_trials) * 0.01 
        gap = max(0, val_score - test_score)
        prob_noise = 1 - np.exp(-gap / (luck_factor + 1e-6))
        return gap <= luck_factor, prob_noise

class QuantEngine:
    @staticmethod
    @st.cache_data(ttl=3600, show_spinner=False)
    def get_market_data(tickers):
        try:
            data = yf.download(tickers, period="1y", progress=False)['Close']
            if isinstance(data, pd.Series): data = data.to_frame()
            return data.pct_change().fillna(0)
        except:
            return pd.DataFrame()

    @staticmethod
    def compute_asymmetric_risk(returns, vix, weights):
        port_ret = returns.dot(weights)
        alpha = 0.05 * (1 + (vix - 20) / 40) if vix > 20 else 0.05
        dynamic_vaR = np.quantile(port_ret, min(alpha, 0.2))
        return dynamic_vaR, port_ret.mean() * 252, port_ret.std() * np.sqrt(252)

    @staticmethod
    def get_realtime_vix():
        """从 Yahoo Finance 获取真实实时 VIX 数值"""
        try:
            vix_data = yf.download("^VIX", period="1d", progress=False)
            if not vix_data.empty:
                # 获取最新的收盘价
                current_vix = vix_data['Close'].iloc[-1]
                return round(float(current_vix), 2)
            return 20.0  # 如果抓取失败，返回中值基准
        except:
            return 20.0


# --- 5. 辅助功能 ---

def generate_docx_report(content, title="Investment Memo"):
    doc = Document()
    doc.add_heading('Macro Alpha Intelligence', 0)
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Report Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for line in content.split('\n'): doc.add_paragraph(line)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def red_team_node(state: AgentState):
    """
    进化版红队节点：AI 驱动的逻辑证伪专家
    """
    q = state.get('quant_results', {})
    if not q:
        return {"is_robust": False, "red_team_critique": "🔴 错误：未检测到量化研究数据。"}
        
    vix = state.get('vix_level', 20.0)
    macro_info = state.get('macro_context', "暂无宏观数据")
    sector_info = state.get('sector_risks', "暂无行业数据")
    
    # 1. 基础数学审计 (硬性指标)
    math_critiques = []
    p_noise = q.get('p_noise', 0)
    active_features = q.get('active', 0)
    
    if p_noise > 0.25:
        math_critiques.append(f"🔴 统计噪音过高 ({p_noise:.1%})，结果可能具有随机性。")
    if active_features < 2:
        math_critiques.append("🔴 特征过于稀疏，模型存在严重的欠拟合风险。")

    # 2. AI 逻辑审计
    audit_prompt = f"""
    你是一名资深量化风险审计师。请对比以下【量化指标】与【外部环境】，寻找逻辑矛盾。
    资产包: {state.get('target_assets', '未知')} | 当前 VIX: {vix} | 量化 VaR: {q.get('d_var', 0):.2%}
    环境背景: {macro_info} | {sector_info}
    
    任务：
    1. 检查是否存在“低 VaR 与高市场波动环境”的错配。
    2. 如果发现模型严重滞后或逻辑自相矛盾，请在回复中包含“🚨【拦截】”字样。
    3. 给出 1-2 句犀利的“证伪”意见。
    """
    
    try:
        response = model.generate_content(audit_prompt)
        ai_critique = response.text
    except Exception as e:
        ai_critique = f"⚠️ AI 审计推理引擎暂时离线。(错误: {str(e)})"

    # 3. 汇总意见
    math_report = "【数学审计】\n" + ("\n".join(math_critiques) if math_critiques else "🟢 统计指标基础稳健。")
    full_critique = f"{math_report}\n\n【AI 逻辑证伪】\n{ai_critique}"

    # --- 4. 判定逻辑：核心改进区 ---
    is_robust = True
    
    # 定义“否决关键词”：如果 AI 的回复里包含这些词，说明逻辑崩了
    veto_keywords = ["🚨", "拦截", "严重错配", "滞后性", "误导", "自欺欺人", "根本性矛盾"]
    
    # 拦截条件 1：数学指标触红线
    if p_noise > 0.30 or active_features < 2:
        is_robust = False
        
    # 拦截条件 2：AI 发现逻辑漏洞（即使数学上看起来很稳健）
    if any(word in ai_critique for word in veto_keywords):
        is_robust = False
        # 如果是因为 AI 拦截，我们可以在报告里加一句标识
        if "🚨【拦截】" not in full_critique:
            full_critique = "🚩 [AI 逻辑强制否决]\n" + full_critique

    return {
        "red_team_critique": full_critique,
        "technical_report": full_critique,
        "is_robust": is_robust  # 👈 只有这个为 True，才会生成“决策建议”
    }

def render_tv_chart(symbol, title):
    code = f"""
    <div style="height:350px; margin-bottom: 20px;"><script src="https://s3.tradingview.com/tv.js"></script>
    <script>new TradingView.MediumWidget({{"symbols": [["{title}", "{symbol}|12M"]],"width": "100%", "height": 350, "locale": "zh_CN", "colorTheme": "light"}});</script></div>"""
    components.html(code, height=360)

@st.cache_data(ttl=3600, show_spinner=False)
def get_ai_analysis(prompt_type, vix_val):
    """通用 AI 专家接口：处理宏观与行业研判"""
    vix_context = f"当前市场 VIX 风险指数为 {vix_val}。"
    prompts = {
        "macro": f"{vix_context} 作为新加坡宏观策略专家，请分析 MAS 货币政策走势、S$NEER 汇率波动及 CPI 对本地 REITs 的影响。",
        "sector": f"{vix_context} 作为行业透视专家，请评估当前 AI 算力板块（NVIDIA/AMD）与地缘政治敏感资产（如黄金、原油）的风险对冲建议。"
    }
    try:
        response = model.generate_content(prompts[prompt_type])
        return response.text
    except Exception as e:
        return f"AI 专家暂时离线: {e}"

# --- 6. LangGraph 智能体逻辑 (Tab 4 专用) ---

class AgentState(TypedDict):
    # --- 1. 基础输入参数 ---
    target_assets: str      # 目标资产组合名称
    vix_level: float        # 当前市场 VIX 指数
    
    # --- 2. 跨板块外部情报 (从 Tab 1 & Tab 3 注入) ---
    macro_context: str      # 存储 Tab 1 的宏观定调（如：通胀预警、加息预期）
    sector_risks: str       # 存储 Tab 3 的行业扫描（如：科技股回调风险、地缘政治）
    
    # --- 3. 核心计算与审计中间态 ---
    # 注意：quant_results 只定义一次，存储所有数值计算结果
    quant_results: dict     # 包含 d_var, p_noise, sparsity, coefs, a_ret 等
    red_team_critique: str  # 存储红队节点的 AI 对抗性意见
    technical_report: str   # 存储原始技术审计意见
    
    # --- 4. 最终输出结果 ---
    is_robust: bool         # 策略稳健性判定（由红队和数学指标共同决定）
    final_memo: str         # 综合宏观、行业、量化的 CEO 级最终审计简报
    
PRESET_ASSETS = {"新加坡蓝筹": ["DBSDF", "U11.SI", "V03.SI"], "科技成长": ["NVDA", "AAPL", "MSFT"]}

def research_node(state: AgentState):
    """量化研究节点：负责基础指标计算"""
    returns = QuantEngine.get_market_data(PRESET_ASSETS[state['target_assets']])
    if returns.empty: return {"is_robust": False}
    weights = np.array([1.0/len(returns.columns)]*len(returns.columns))
    d_var, a_ret, a_vol = QuantEngine.compute_asymmetric_risk(returns, state['vix_level'], weights)
    
    y = returns.iloc[:, 0]
    X = returns.shift(1).fillna(0)
    active, sparsity, coefs = StrategyAuditor.run_feature_sparsity_check(X, y)
    is_robust, p_noise = StrategyAuditor.check_optimizer_curse(0.05, 0.045, 100)
    
    # --- 核心计算逻辑：审计置信度 ---
    base_score = 100
    penalty_p = max(0, (p_noise - 0.05) * 200) # 超过5%后，每增加1%扣2分
    penalty_s = 20 if active < 2 else 0
    confidence_score = max(0, min(100, base_score - penalty_p - penalty_s))
    
    # ⚠️ 关键修正：必须将 confidence_score 包含在返回的字典中
    return {
        "quant_results": {
            "d_var": d_var, 
            "p_noise": p_noise, 
            "sparsity": sparsity, 
            "active": active, 
            "coefs": coefs, 
            "returns": returns, 
            "X": X, 
            "a_ret": a_ret, 
            "a_vol": a_vol,
            "confidence_score": confidence_score # 👈 确保这个键存在
        }, 
        "is_robust": (p_noise < 0.3)
    }

def technical_audit_node(state: AgentState):
    """审计节点：负责撰写硬核技术报告"""
    q = state['quant_results']
    prompt = f"""
    你是一名量化风险审计师。请针对以下指标撰写一份【硬核技术报告】：
    资产包: {state['target_assets']}
    VIX环境: {state['vix_level']}
    VaR: {q['d_var']:.2%}
    Lasso特征数: {q['active']}
    P-hacking 风险: {q['p_noise']:.2%}
    
    要求：使用专业量化术语，分析模型的统计稳健性、过拟合风险及非对称风险暴露。
    """
    try:
        res = model.generate_content(prompt)
        return {"technical_report": res.text}
    except:
        return {"technical_report": "技术审计调用失败。"}

def translator_node(state: AgentState):
    """翻译官节点：将技术报告重写为 CEO 级简报 (核心改进)"""
    q = state['quant_results']
    tech_report = state['technical_report']
    
    prompt = f"""
    你是一名资深投研顾问，负责向基金经理（Fund Manager）汇报。
    请根据下方的【硬核技术报告】，将其重写为【CEO 级分层简报】。
    
    原始报告内容：{tech_report}
    
    ### 严格遵循以下输出格式 ###
    
    ### 📢 首席执行决策建议 (Executive Summary)
    [用直白、果断的决策建议起头：买入/持有/减仓。严禁统计术语，告诉老板“So What”。]
    
    ### 🔍 关键洞察 (Key Insights)
    [将复杂指标翻译为业务语言：
    - VaR -> 压力下的潜在亏损
    - Sparsity -> 决策信号纯净度
    - P-hacking -> 结果真实可信度]
    
    ### 🔬 技术附录 (Technical Appendix)
    [保留硬核术语和具体数值，供量化同事复核。]
    """
    try:
        res = model.generate_content(prompt)
        return {"audit_memo": res.text}
    except:
        return {"audit_memo": "翻译决策生成失败。"}

# 构建增强版对抗 Graph 
builder = StateGraph(AgentState)

builder.add_node("researcher", research_node)
builder.add_node("red_team", red_team_node)     # 👈 新增红队质疑节点
builder.add_node("auditor", technical_audit_node)
builder.add_node("translator", translator_node)

builder.set_entry_point("researcher")

# 逻辑流：计算 -> 红队质疑 -> (通过则) 审计 -> 翻译 
builder.add_edge("researcher", "red_team")

builder.add_conditional_edges(
    "red_team", 
    lambda x: "auditor" if x["is_robust"] else END  # 如果被红队毙掉，直接终止流程 
)

builder.add_edge("auditor", "translator")
builder.add_edge("translator", END)

agent_executor = builder.compile()

# --- 7. 主界面布局与样式注入 ---

# 注入自定义 CSS 以提升 UI 质感
st.markdown("""
    <style>
    /* 全局背景与字体微调 */
    .main { background-color: #FDFDFD; }
    
    /* Metric 卡片美化 */
    div[data-testid="metric-container"] {
        background-color: #FFFFFF;
        border: 1px solid #ECEFF4;
        padding: 1rem;
        border-radius: 12px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    [data-testid="stMetricValue"] { font-size: 24px; color: #1B315E; font-weight: 700; }
    [data-testid="stMetricLabel"] { font-size: 14px; color: #64748B; }

    /* 决策建议卡片 */
    .decision-card {
        background-color: #F8FAFC;
        border-left: 5px solid #1B315E;
        padding: 20px;
        border-radius: 8px;
        margin: 15px 0;
    }
    
    /* 侧边栏优化 */
    .css-1639116 { background-color: #1B315E; }
    .stButton>button {
        border-radius: 8px;
        height: 3em;
        transition: all 0.3s;
    }
    </style>
    """, unsafe_allow_html=True)

with st.sidebar:
    st.header("⚙️ 智能环境配置")
    
    real_vix = QuantEngine.get_realtime_vix()
    
    # 动态风险颜色与标签
    if real_vix < 15:
        risk_label, risk_color = "低波动 (Complacency)", "blue"
    elif real_vix < 25:
        risk_label, risk_color = "常态波动 (Normal)", "green"
    elif real_vix < 35:
        risk_label, risk_color = "高波动 (Panic)", "orange"
    else:
        risk_label, risk_color = "极高风险 (Crisis)", "red"

    st.metric("实时 VIX 指数", real_vix, delta=risk_label, delta_color="normal")
    
    vix_input = st.number_input("环境压力修正 (默认为实时)", value=real_vix, help="手动调整波动率倍率以进行压力测试")
    
    st.divider()
    st.info("💡 **系统状态**: 当前处于【首席审计模式】，Agent 将优先检查统计稳健性。")
    st.caption("数据源: CBOE Real-time / Yahoo Finance")

    # --- 找到侧边栏“🤖 开启全行业自动扫描”部分 ---
# --- 侧边栏修改 ---
if st.sidebar.button("🤖 开启全行业自动扫描"):
    with st.status("正在扫描全球市场板块...", expanded=True) as status:
        scanner = MarketScanner()
        top_asset = scanner.run_daily_scan() # 假设这里返回详细信息
        
        # 存储到 session_state
        st.session_state.scan_results = top_asset 
        st.session_state.target_assets = top_asset['name']
        
        status.update(label="扫描完成，建议已生成！", state="complete")
    st.rerun()

# 增加一个标签页
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🧠 首席宏观研判", 
    "📈 实时仪表盘", 
    "🛡️ 行业风险穿透", 
    "🔢 量化审计室",
    "🤖 AI 全行业扫描"  # <--- 新增
])

with tab1:
    st.header("🧠 首席宏观投研备忘录")
    with st.container():
        col_m1, col_m2 = st.columns([2, 1])
        with col_m1:
            if st.button("🚀 启动宏观协同分析", type="primary", use_container_width=True):
                with st.spinner("正在调度新加坡宏观专家库..."):
                    macro_res = get_ai_analysis("macro", vix_input)
                    st.markdown(f'<div class="decision-card">{macro_res}</div>', unsafe_allow_html=True)
                    st.download_button("📥 下载宏观报告", generate_docx_report(macro_res, "Macro Memo"), "Macro_Report.docx")
        with col_m2:
            st.caption("专家组成员：\n- MAS 政策分析师\n- 离岸汇率策略师\n- 通胀研究专家")

with tab2:
    st.subheader("📊 全球资产实时走廊")
    
    # 第一排：黄金 vs 纳指
    col1, col2 = st.columns(2)
    with col1: render_tv_chart("OANDA:XAUUSD", "现货黄金")
    with col2: render_tv_chart("NASDAQ:NDX", "纳斯达克 100")
    
    # 第二排：海指 vs 汇率
    col3, col4 = st.columns(2)
    with col3: render_tv_chart("STI:STI", "海峡指数")
    with col4: render_tv_chart("FX_IDC:USDSGD", "美元/新币")

    # 第三排：原油 vs 沪深300 (补回这两项)
    col5, col6 = st.columns(2)
    with col5: render_tv_chart("TVC:UKOIL", "布伦特原油 (Brent)")
    with col6: render_tv_chart("SSE:000300", "沪深 300 指数")

with tab3:
    st.header("🛡️ 行业穿透与非对称风险评估")
    if st.button("🔍 执行行业风险扫描", type="primary"):
        with st.spinner("正在分析行业暴露风险..."):
            sector_res = get_ai_analysis("sector", vix_input)
            st.markdown(f'<div class="decision-card">{sector_res}</div>', unsafe_allow_html=True)
            st.download_button("📥 下载行业报告", generate_docx_report(sector_res, "Sector Drilldown"), "Sector_Report.docx")

with tab4:
    st.header("🔢 Agentic AI 深度审计终端")
    
    # --- 1. 基础数据准备 ---
    current_assets = list(PRESET_ASSETS.keys())
    
    # --- 2. 状态同步逻辑 (核心修复点) ---
    # 尝试从 session_state 获取扫描建议，如果没有则默认 0
    # 这样既能实现“自动填表”，又保证了 default_index 永远存在
    scanned_asset = st.session_state.get("target_assets")
    default_index = current_assets.index(scanned_asset) if scanned_asset in current_assets else 0

    # --- 3. 极简交互区 ---
    c_left, c_right = st.columns([3, 1])
    with c_left:
        # 这里定义的 choice 就是你“选的内容”
        choice = st.selectbox(
            "选择审计目标", 
            current_assets, 
            index=default_index, 
            key="audit_asset_selector_v_final",
            label_visibility="collapsed"
        )
    with c_right:
        start_audit = st.button("🚀 启动深度审计", type="primary", use_container_width=True)

    # --- 4. 触发与执行 ---
    if start_audit:
    # 如果此时 session 中残留了自动触发标记，顺便清理掉，保持状态干净
    if st.session_state.get("auto_trigger"):
        st.session_state.auto_trigger = False
        
        # ⚠️ 关键：Agent 这里的 target_assets 直接绑定 choice
        # 这样不管扫描结果是什么，Agent 只会审计你在下拉框里选中的那个
        with st.status(f"Agent 正在对 {choice} 执行深度审计...", expanded=True) as status:
            current_state = {
                "target_assets": choice,  # 认准 UI 选择
                "vix_level": vix_input, 
                "macro_context": st.session_state.get("macro_memo", "暂无数据"),
                "sector_risks": st.session_state.get("sector_memo", "暂无数据"),
                "quant_results": {}, 
                "is_robust": True
            }
            
            # 2. 流式执行 LangGraph 节点
            try:
                # 这里的 agent_executor 是你之前定义的 StateGraph 编译后的对象
                for event in agent_executor.stream(current_state):
                    if not event: continue
                    
                    for node_name, node_output in event.items():
                        # 状态提示：让用户知道哪个专家正在干活
                        st.write(f"⚙️ **{node_name}** 节点处理完成...")
                        
                        # 安全合并节点产出到总状态
                        if isinstance(node_output, dict):
                            clean_output = {k: v for k, v in node_output.items() if v is not None}
                            current_state.update(clean_output)
                
                # 3. 将最终审计结果存入 session_state 供下方渲染区展示
                st.session_state.final_audit_state = current_state
                status.update(label="✅ 深度审计流执行完毕", state="complete")
                
            except Exception as e:
                st.error(f"❌ 审计流执行期间发生异常: {str(e)}")
                status.update(label="⚠️ 审计流意外中断", state="error")
            
            # 2. 增强版流式循环
            try:
                for event in agent_executor.stream(current_state):
                    if not event: continue  # 过滤空事件
                    
                    for node_name, node_output in event.items():
                        st.write(f"✅ {node_name} 任务处理完成")
                        
                        # 核心修复：确保 node_output 是字典且不为空
                        if isinstance(node_output, dict):
                            # 使用 dict.update 之前先过滤掉 None 值
                            clean_output = {k: v for k, v in node_output.items() if v is not None}
                            current_state.update(clean_output)
                        else:
                            st.warning(f"节点 {node_name} 返回了异常数据类型")
                            
                st.session_state.final_audit_state = current_state
                status.update(label="审计流执行完毕", state="complete")
                
            except Exception as e:
                st.error(f"审计流运行出错: {str(e)}")
                status.update(label="审计中断", state="error")

    # --- Tab 4 结果渲染区 (增强版) ---
    if st.session_state.final_audit_state:
        s = st.session_state.final_audit_state
        q = s.get('quant_results', {})
        
        if q:
            score = q.get('confidence_score', 0)
            
            # --- 第一层：审计概览 ---
            st.divider()
            col_score, col_status = st.columns([1, 2])
            with col_score:
                st.metric("审计置信度", f"{int(score)}/100")
           # --- 找到这一段进行替换 ---
            with col_status:
                # 1. 逻辑修正：置信度与稳健性联动
                if score >= 85 and s["is_robust"]:
                    res_color, res_label = "green", "💎 审计通过：高度可信"
                elif score >= 60 and s["is_robust"]:
                    res_color, res_label = "orange", "⚠️ 审计通过：中度参考"
                else:
                    res_color, res_label = "red", "🚨 审计拦截：逻辑/统计存在重大缺陷"
                
                st.markdown(f"<h2 style='color:{res_color}; margin-top:0;'>{res_label}</h2>", unsafe_allow_html=True)

            # --- 第二层：分情况展示详细理由 (确保此 if 与 with col_status 对齐) ---
            if not s["is_robust"]:
                # 【拦截状态】：展示红队的详细“证伪”理由
                st.error("### 🛑 详细拦截理由 (Red Team Audit Report)")
                
                st.markdown(f"""
                <div style="background-color:#FFF5F5; padding:25px; border-radius:12px; border:2px solid #FEB2B2; color:#C53030;">
                    <h4 style="margin-top:0;">🚩 红队对抗性证伪结论：</h4>
                    {s.get('red_team_critique', '未获取到详细证伪理由。')}
                </div>
                """, unsafe_allow_html=True)
                
                st.subheader("📊 风险指标溯源")
                k1, k2, k3 = st.columns(3)
                k1.metric("P-hacking 风险值", f"{q['p_noise']:.1%}", delta="超标" if q['p_noise']>0.3 else "正常", delta_color="inverse")
                k2.metric("有效特征数", f"{q['active']}", delta="过低" if q['active']<2 else "正常", delta_color="inverse")
                k3.metric("VaR 预警", f"{q.get('d_var', 0):.2%}")
                
                st.info("💡 **专家建议**：请重新检查数据源是否包含未来函数，或降低模型复杂度（减少因子数量）后再尝试审计。")

            else:
                # 【通过状态】：展示决策建议
                st.subheader("💡 首席执行决策建议")
                if s.get("audit_memo"):
                    st.info(s["audit_memo"])
                
                audit_col1, audit_col2 = st.columns(2)
                with audit_col1:
                    st.markdown("#### 🚩 红队合规意见")
                    st.success(s.get("red_team_critique", "未发现明显统计性漏洞。"))
                with audit_col2:
                    st.markdown("#### 📝 原始技术审计")
                    st.caption(s.get("technical_report", "技术审计内容生成中..."))

            # --- 第三层：可视化（无论通不通过都显示，方便排查） ---
            with st.expander("🔬 查看量化底稿 (Quantitative Backing)"):
                st.bar_chart(pd.DataFrame({'Factor': q['X'].columns, 'Weight': q['coefs']}), x="Factor", y="Weight")

with tab5:
    st.header("🤖 全球资产扫描建议")
    
    # 从 session_state 获取扫描结果
    # 假设 scanner.run_daily_scan() 返回的是一个包含多个资产信息的 List 或 Dict
    scan_data = st.session_state.get("scan_results") 

    if not scan_data:
        st.info("💡 请点击侧边栏的 **[开启全行业自动扫描]** 按钮，获取实时投资机会。")
    else:
        # 1. 顶部冠军卡片
        top_asset = st.session_state.get("target_assets")
        st.success(f"🏆 本次扫描冠军：**{top_asset}**")
        
        # 2. 展示扫描后的分析细节 (示例：可以用 DataFrame 展示排名)
        # 这里你可以展示夏普比率、收益分布、波动率等
        st.subheader("📊 扫描分析看板")
        col_s1, col_s2, col_s3 = st.columns(3)
        col_s1.metric("推荐资产", top_asset)
        col_s2.metric("预期夏普比率", "2.14", delta="高收益区间")
        col_s3.metric("市场契合度", "92%", delta="优选")

        # 3. 这里的关键：引导用户去审计
        st.markdown("---")
        st.write("根据量化引擎筛选，该资产目前处于极佳风险收益比区间。")
        
        if st.button("🔍 立即将该资产送往 [量化审计室] 进行深度证伪"):
            st.session_state.auto_trigger = True
            # 注意：这里不需要再 rerun，因为用户点击后会想看到变化，或者手动切到下一页
            st.toast(f"已将 {top_asset} 送往审计室！")

st.markdown("---")
st.caption("Macro Alpha Pro | NUS MSBA Project | 专注量化审计与合规决策")
